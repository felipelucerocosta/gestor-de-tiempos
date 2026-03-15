from openpyxl import Workbook
from docx import Document
import database as db
import logic
import os

def export_to_excel(project_id, file_path):
    projects = db.get_projects()
    project = next(p for p in projects if p[0] == project_id)
    logs = db.get_logs(project_id)
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Reporte de Proyecto"
    
    # Header
    ws.append(["Proyecto", project[1]])
    ws.append(["Fecha Límite", project[3]])
    ws.append([])
    
    # Logs
    ws.append(["Fecha", "Nivel", "Descripción"])
    for l in logs:
        level_text = "Menor" if l[4] == 1 else "Mayor" if l[4] == 5 else "Ninguno"
        ws.append([l[2], level_text, l[3]])
    
    wb.save(file_path)

def export_to_word(project_id, file_path):
    projects = db.get_projects()
    project = next(p for p in projects if p[0] == project_id)
    logs = db.get_logs(project_id)
    
    doc = Document()
    doc.add_heading(f"Reporte de Proyecto: {project[1]}", 0)
    
    doc.add_paragraph(f"Fecha de Creación: {project[2]}")
    doc.add_paragraph(f"Fecha Límite: {project[3]}")
    
    doc.add_heading("Avances del Proyecto", level=1)
    
    for l in logs:
        level_text = "Menor" if l[4] == 1 else "Mayor" if l[4] == 5 else "Ninguno"
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"[{l[2][:10]}] - {level_text}: ").bold = True
        p.add_run(l[3])
    
    doc.save(file_path)
